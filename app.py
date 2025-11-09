from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4, LETTER, LEGAL
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import traceback
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
import json

app = Flask(__name__)
CORS(app)


# ===== DOCX Helper Functions =====
def apply_docx_attributes(run, attrs):
    """Apply Quill attributes to a docx run."""
    if not attrs:
        return
    if attrs.get("bold"):
        run.bold = True
    if attrs.get("italic"):
        run.italic = True
    if attrs.get("underline"):
        run.underline = True
    if attrs.get("strike"): 
        run.font.strike = True
    # 🎨 Handle color (supports both #RRGGBB and #AARRGGBB)
    if attrs.get("color"):
        try:
            hex_color = attrs["color"].lstrip("#")

            # If alpha channel is present (e.g., #FFE53935), strip it
            if len(hex_color) == 8:
                hex_color = hex_color[2:]  # remove first 2 (alpha)

            if len(hex_color) == 6:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            else:
                print(f"⚠️ Unexpected color format: {attrs['color']}")

        except Exception as e:
            print("⚠️ Invalid color:", e)


def set_paragraph_alignment_docx(para, attrs):
    """Set paragraph alignment."""
    if not attrs:
        return
    align = attrs.get("align")
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def create_list_paragraph(doc, text, attrs, list_type="bullet"):
    """Create a properly aligned list paragraph with smaller indent."""
    para = doc.add_paragraph(text, style="List Bullet" if list_type == "bullet" else "List Number")
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.25)
    para_format.space_after = Pt(0)
    para_format.line_spacing = Pt(14)
    set_paragraph_alignment_docx(para, attrs)
    for run in para.runs:
        apply_docx_attributes(run, attrs)
    return para


# ===== MAIN DOCX CONVERSION =====
@app.route("/convert/delta-to-docx", methods=["POST"])
def delta_to_docx():
    try:
        data = request.get_json()
        delta = data.get("delta")
        page_size_name = str(data.get("page_size", "A4")).strip().lower()
        margins = data.get("margins", {"top": 20, "bottom": 20, "left": 20, "right": 20})

        if not delta:
            return jsonify({"error": "No delta provided"}), 400

        doc = Document()

        # === PAGE SETUP ===
        section = doc.sections[0]
        if "letter" in page_size_name:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
        elif "legal" in page_size_name:
            section.page_height = Inches(14)
            section.page_width = Inches(8.5)
        else:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)

        mm_to_inch = lambda mm: mm / 25.4
        section.top_margin = Inches(mm_to_inch(margins["top"]))
        section.bottom_margin = Inches(mm_to_inch(margins["bottom"]))
        section.left_margin = Inches(mm_to_inch(margins["left"]))
        section.right_margin = Inches(mm_to_inch(margins["right"]))

        # === PROCESS DELTA ===
        current_para = None
        list_mode = None

        for op in delta:
            insert_data = op.get("insert")
            attrs = op.get("attributes", {})

            # ---- Handle custom tables (Firebase) ----
            if isinstance(insert_data, dict) and "custom" in insert_data:
                try:
                    custom_data = json.loads(insert_data["custom"])
                    table_json_str = custom_data.get("table")
                    if table_json_str:
                        table_data = json.loads(table_json_str)
                        rows = table_data.get("rows", [])
                        columns = table_data.get("columns", [])
                        if rows:
                            table = doc.add_table(rows=len(rows), cols=len(columns))
                            for r_idx, row in enumerate(rows):
                                for c_idx, col_key in enumerate(columns):
                                    cell_ops = row.get(col_key, [])
                                    cell_para = table.cell(r_idx, c_idx).paragraphs[0]
                                    for cell_op in cell_ops:
                                        text = cell_op.get("insert", "")
                                        run = cell_para.add_run(text)
                                        apply_docx_attributes(run, cell_op.get("attributes", {}))
                            current_para = None
                            continue
                except Exception as e:
                    print("⚠️ Table parse error:", e)
                    continue

            # ---- Plain text or list ----
            if isinstance(insert_data, str):
                list_type = attrs.get("list")

                # Start new paragraph if necessary
                if current_para is None or (list_type and list_mode != list_type):
                    style = "List Bullet" if list_type == "bullet" else "List Number" if list_type == "ordered" else None
                    current_para = doc.add_paragraph(style=style)
                    list_mode = list_type

                    pf = current_para.paragraph_format
                    pf.left_indent = Inches(0.25 if list_type else 0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = Pt(12)

                # Handle line breaks within insert
                parts = insert_data.split("\n")
                for i, part in enumerate(parts):
                    if part:
                        run = current_para.add_run(part)
                        apply_docx_attributes(run, attrs)
                    if i < len(parts) - 1:  # newline means paragraph end
                        set_paragraph_alignment_docx(current_para, attrs)
                        current_para = doc.add_paragraph(style=("List Bullet" if list_type == "bullet" else "List Number" if list_type == "ordered" else None))
                        pf = current_para.paragraph_format
                        pf.left_indent = Inches(0.25 if list_type else 0)
                        pf.space_after = Pt(0)
                        pf.line_spacing = Pt(12)

                # Apply alignment at the end of this op
                set_paragraph_alignment_docx(current_para, attrs)

        # === SAVE DOCX ===
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name="document.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("❌ Error:", traceback.format_exc())
        return jsonify({"error": str(e)}), 500



# ===== PDF Helper Functions =====
def get_pdf_alignment(align):
    if align == "center":
        return TA_CENTER
    elif align == "right":
        return TA_RIGHT
    elif align == "justify":
        return TA_JUSTIFY
    return TA_LEFT

def parse_color(color_code):
    """Convert hex color to ReportLab color."""
    try:
        if not color_code.startswith("#"):
            color_code = f"#{color_code}"
        return colors.HexColor(color_code)
    except:
        return colors.black

def apply_text_styles(text, attrs):
    """Apply all inline styles: bold, italic, underline, strike, color, background."""
    if not text:
        return ""
    if not attrs:
        return text.replace("\n", "<br/>")
    
    styled_text = text
    if attrs.get("bold"):
        styled_text = f"<b>{styled_text}</b>"
    if attrs.get("italic"):
        styled_text = f"<i>{styled_text}</i>"
    if attrs.get("underline"):
        styled_text = f"<u>{styled_text}</u>"
    if attrs.get("strike"):
        styled_text = f"<strike>{styled_text}</strike>"
    if attrs.get("color"):
        styled_text = f'<font color="{attrs["color"]}">{styled_text}</font>'
    if attrs.get("background"):
        styled_text = f'<font backColor="{attrs["background"]}">{styled_text}</font>'
    return styled_text.replace("\n", "<br/>")


# ===== PDF Conversion =====
@app.route("/convert/delta-to-pdf", methods=["POST"])
def delta_to_pdf():
    try:
        import json
        data = request.get_json()
        delta = data.get("delta")
        page_size_name = str(data.get("page_size", "A4")).strip().lower()
        margins = data.get("margins", {"top": 20, "bottom": 20, "left": 20, "right": 20})

        if not delta:
            return jsonify({"error": "No delta provided"}), 400

        PAGE_SIZES = {"a4": A4, "letter": LETTER, "legal": LEGAL}
        page_size = PAGE_SIZES.get(page_size_name, A4)

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            topMargin=margins["top"],
            bottomMargin=margins["bottom"],
            leftMargin=margins["left"],
            rightMargin=margins["right"]
        )

        story = []
        styles = getSampleStyleSheet()

        paragraph_ops = []
        list_blocks = []
        current_alignment = None

        def flush_paragraph():
            """Render one full paragraph or list item."""
            nonlocal paragraph_ops, current_alignment
            if not paragraph_ops:
                return

            # Merge all styled fragments
            html_text = "".join([
                apply_text_styles(op["insert"], op.get("attributes", {}))
                for op in paragraph_ops
                if isinstance(op.get("insert"), str)
            ]).strip()

            if not html_text:
                paragraph_ops = []
                current_alignment = None
                return

            paragraph_style = ParagraphStyle(
                name="Custom",
                parent=styles["Normal"],
                alignment=get_pdf_alignment(current_alignment),
                leading=14,
                spaceAfter=6,
            )

            if current_list_type == "bullet":
                story.append(Paragraph(html_text, paragraph_style, bulletText="•"))
            elif current_list_type == "ordered":
                story.append(Paragraph(f"{list_counters['ordered']}. {html_text}", paragraph_style))
                list_counters["ordered"] += 1
            else:
                story.append(Paragraph(html_text, paragraph_style))

            story.append(Spacer(1, 4))
            paragraph_ops = []
            current_alignment = None

        current_list_type = None
        list_counters = {"ordered": 1}

        for op in delta:
            insert_data = op.get("insert")
            attrs = op.get("attributes", {})

            # Handle table
            if isinstance(insert_data, dict) and "custom" in insert_data:
                try:
                    custom_data = json.loads(insert_data["custom"])
                    table_json_str = custom_data.get("table")
                    if table_json_str:
                        table_data_parsed = json.loads(table_json_str)
                        rows = table_data_parsed.get("rows", [])
                        col_keys = table_data_parsed.get("columns", [])
                        col_count = len(col_keys)
                        data_rows = []

                        for row in rows:
                            row_cells = []
                            for col_key in col_keys:
                                cell_ops = row.get(col_key, [])
                                cell_text = ""
                                for cell_op in cell_ops:
                                    cell_text += apply_text_styles(cell_op.get("insert", ""), cell_op.get("attributes", {}))
                                row_cells.append(Paragraph(cell_text or "&nbsp;", styles["Normal"]))
                            data_rows.append(row_cells)

                        table = Table(
                            data_rows,
                            colWidths=[(page_size[0] - margins["left"] - margins["right"]) / col_count] * col_count,
                            style=TableStyle([
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ])
                        )
                        story.append(table)
                        story.append(Spacer(1, 12))
                        continue
                except Exception as e:
                    print("⚠️ Table parse error:", e)
                    continue

            # Handle text
            if isinstance(insert_data, str):
                parts = insert_data.split("\n")
                for i, part in enumerate(parts):
                    if part:
                        paragraph_ops.append({"insert": part, "attributes": attrs})

                    # If newline marks the end of a paragraph or list
                    if i < len(parts) - 1:
                        # Capture alignment and list from newline attributes
                        current_alignment = attrs.get("align", current_alignment)
                        list_attr = attrs.get("list")

                        if list_attr:
                            current_list_type = list_attr
                        else:
                            current_list_type = None

                        flush_paragraph()


        flush_paragraph()

        doc.build(story)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.pdf",
            mimetype="application/pdf"
        )

    except Exception as e:
        print("❌ Error:", traceback.format_exc())
        return jsonify({"error": str(e)}), 500


@app.route("/routes")
def list_routes():
    return jsonify([str(rule) for rule in app.url_map.iter_rules()])

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
